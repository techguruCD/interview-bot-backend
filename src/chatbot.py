from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import Pinecone
from langchain.llms.base import LLM
from langchain.chains.question_answering import load_qa_chain
from langchain.utils import get_from_dict_or_env
import together
import docx
from langchain.schema import Document
import os
from constants import *
from config import *
import pickle
from pydantic import Extra, root_validator
from typing import Any, Dict
import pinecone
import PyPDF2
import traceback
import openai

openai.api_type = API_TYPE
openai.api_key = OPENAI_API_KEY
openai.api_base = ENDPOINT
openai.api_version = API_VERSION


class TogetherLLM(LLM):
    """Together large language models."""

    model: str = "togethercomputer/llama-2-7b-chat"
    """model endpoint to use"""

    together_api_key: str = os.environ["TOGETHER_API_KEY"]
    """Together API key"""

    temperature: float = TEMPRATURE
    """What sampling temperature to use."""
    top_p: float = 0.3

    max_tokens: int = 300
    """The maximum number of tokens to generate in the completion."""

    class Config:
        extra = Extra.forbid

    @root_validator()
    def validate_environment(cls, values: Dict) -> Dict:
        """Validate that the API key is set."""
        api_key = get_from_dict_or_env(values, "together_api_key", "TOGETHER_API_KEY")
        values["together_api_key"] = api_key
        return values

    @property
    def _llm_type(self) -> str:
        """Return type of LLM."""
        return "together"

    def _call(
        self,
        prompt: str,
        **kwargs: Any,
    ) -> str:
        """Call to Together endpoint."""
        together.api_key = self.together_api_key
        output = together.Complete.create(
            prompt,
            model=self.model,
            max_tokens=self.max_tokens,
            temperature=self.temperature,
        )
        text = output["output"]["choices"][0]["text"]
        return text


def load_vector_store():
    with open(EMBEDDING_FILE, "rb") as f:
        vectorstore = pickle.load(f)
    return vectorstore


def initialize_qa_chain(LLM_OBJECT, prompt=PROMPT):
    qa = load_qa_chain(LLM_OBJECT, chain_type="stuff", prompt=prompt)

    return qa


def get_llm_object():
    together.api_key = os.environ["TOGETHER_API_KEY"]
    together.Models.start("togethercomputer/llama-2-7b-chat")
    llm = TogetherLLM(
        model="togethercomputer/llama-2-7b-chat",
        temperature=TEMPRATURE,
        max_tokens=MAX_TOKEN,
    )
    return llm


def get_response(query, chain, doc):
    response = chain.run(input_documents=doc, question=query)
    return response


def create_document(pdf, extra_info, user_id, logger):
    try:
        if pdf.lower().endswith(".docx"):
            doc = docx.Document(pdf)
            text = []
            for para in doc.paragraphs:
                text.append(para.text)
            text = "\t".join(text)
        elif pdf.lower().endswith(".pdf"):
            with open(pdf, "rb") as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                text = ""
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text()
        else:
            text = ""
        text = text + extra_info

        documents = []
        documents.append(Document(page_content=text, metadata={"id": user_id}))
    except Exception as e:
        logger.exception(
            f"Error in create_document, \n TRACEBACK:{traceback.format_exc()}",
            exc_info=True,
        )
    return documents


def get_embedding():
    embeddings = OpenAIEmbeddings(
        deployment=DEPLOYMENT_NAME,
        model="text-embedding-ada-002",
        openai_api_base=ENDPOINT,
        openai_api_type="azure",
        openai_api_key=OPENAI_API_KEY,
    )
    return embeddings


def prompt(tone, style, name):
    txt = "use the {} tone and style of conversation should be {}".format(tone, style)
    inst = INSTRUCTIONS.format(name)
    prompt_ = inst + txt + TEMPLATE
    return prompt_


def save_pickel(file):
    try:
        with open(EMBEDDING_FILE, "wb") as f:
            pickle.dump(file, f)
    except Exception as e:
        return e
    return True


def save_embedding(doc, user_id, embedding, index, logger):
    try:
        pinecone_db = Pinecone.from_documents(
            doc, embedding, index_name=index, ids=user_id
        )
        logger.info("Embedding saved successfully: ")
        return True
    except Exception as e:
        logger.exception(
            f"Error while saving embedding:, \n TRACEBACK:{traceback.format_exc()}",
            exc_info=True,
        )

        return False


def get_embedding_pinecone(user_id, index, embedding, logger):
    try:
        pinecone_db = Pinecone.from_existing_index(index, embedding)
        query = ""
        response_from_pinecone = pinecone_db.similarity_search(
            query, filter={"id": {"$eq": user_id}}
        )
        if response_from_pinecone != " ":
            return True, response_from_pinecone
        else:
            return False, "no data found in pinecone"

    except pinecone.ApiException as e:
        logger.exception(
            f"Error while retrieving embedding:, \n TRACEBACK:{traceback.format_exc()}",
            exc_info=True,
        )
        return False, f"Error while retrieving embedding: {e}"
